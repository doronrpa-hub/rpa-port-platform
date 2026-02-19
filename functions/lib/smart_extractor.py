"""
Multi-method document extractor for RCB.

RULES:
1. NEVER return "no data" until ALL applicable methods have been tried.
2. When multiple methods return data, COMPARE and pick the best.
3. When methods disagree on critical fields, FLAG for review.
4. ALWAYS explain what was tried and what happened.

Session 28C — Assignment 20.
NEW FILE — does not modify any existing code.
"""

import io
import re
import csv
import logging
from .extraction_result import ExtractionResult

logger = logging.getLogger("rcb.smart_extractor")


class SmartExtractor:
    """Multi-method document extractor. Tries everything, picks the best."""

    def extract(self, file_bytes, filename, content_type):
        """
        Master extraction. Tries all applicable methods, picks best result.

        Args:
            file_bytes: bytes — raw file content
            filename: str — original filename
            content_type: str — MIME type

        Returns:
            ExtractionResult
        """
        if not file_bytes:
            return ExtractionResult(
                text="", confidence=0.0, method="none", is_empty=True,
                warnings=["Empty file provided (0 bytes)."],
            )

        results = []
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        ct = (content_type or "").lower()

        # Route to format-specific extractors
        if ct == "application/pdf" or ext == "pdf":
            results = self._extract_pdf(file_bytes)
        elif ct in ("text/html", "application/xhtml+xml") or ext in ("html", "htm"):
            results = self._extract_html(file_bytes)
        elif "spreadsheet" in ct or "excel" in ct or ext in ("xlsx", "xls"):
            results = self._extract_excel(file_bytes)
        elif "wordprocessingml" in ct or ext in ("docx", "doc"):
            results = self._extract_word(file_bytes)
        elif ct.startswith("image/") or ext in ("png", "jpg", "jpeg", "tiff", "tif", "bmp", "gif"):
            results = self._extract_image(file_bytes)
        elif ct in ("text/csv", "text/tab-separated-values") or ext in ("csv", "tsv"):
            results = self._extract_csv(file_bytes)

        # Generic fallbacks if nothing matched or all failed
        if not results:
            r = self._try_plain_text(file_bytes)
            if r:
                results.append(r)

        if not results:
            r = self._try_vision_ocr_fallback(file_bytes, filename)
            if r:
                results.append(r)

        if not results:
            return ExtractionResult(
                text="", confidence=0.0, method="none", is_empty=True,
                warnings=[
                    f"ALL extraction methods failed for '{filename}' ({content_type}). "
                    f"Document may be encrypted, corrupted, empty, or unsupported."
                ],
            )

        return self._pick_best(results)

    # ═══════════════════════════════════════════
    #  PDF extraction — 4 methods
    # ═══════════════════════════════════════════

    def _extract_pdf(self, file_bytes):
        results = []

        # Method 1: pdfplumber — best for text PDFs with tables
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages_text = []
                tables = []
                for page in pdf.pages[:50]:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                    page_tables = page.extract_tables()
                    if page_tables:
                        for t in page_tables:
                            parsed = self._table_to_dicts(t)
                            if parsed:
                                tables.append(parsed)

                full_text = "\n".join(pages_text)
                if full_text.strip():
                    confidence = min(0.9, 0.5 + len(full_text) / 5000)
                    results.append(ExtractionResult(
                        text=full_text, confidence=confidence,
                        method="pdfplumber", tables=tables,
                    ))
        except Exception as e:
            logger.debug(f"pdfplumber failed: {e}")

        # Method 2: pypdf — different parser, catches some pdfplumber misses
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages[:50]:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            full_text = "\n".join(pages_text)
            if full_text.strip():
                confidence = min(0.8, 0.4 + len(full_text) / 5000)
                results.append(ExtractionResult(
                    text=full_text, confidence=confidence, method="pypdf",
                ))
        except Exception as e:
            logger.debug(f"pypdf failed: {e}")

        # Method 3: PyMuPDF (fitz) — yet another parser
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages_text = []
            for page_num in range(min(len(doc), 50)):
                text = doc[page_num].get_text()
                if text:
                    pages_text.append(text)
            doc.close()
            full_text = "\n".join(pages_text)
            if full_text.strip():
                confidence = min(0.85, 0.45 + len(full_text) / 5000)
                results.append(ExtractionResult(
                    text=full_text, confidence=confidence, method="pymupdf",
                ))
        except Exception as e:
            logger.debug(f"PyMuPDF text extraction failed: {e}")

        # Method 4: Vision OCR — only if text methods failed or low confidence
        # This costs money so we only try when needed.
        if not results or all(r.confidence < 0.5 for r in results):
            try:
                ocr_text = self._vision_ocr_pdf(file_bytes)
                if ocr_text and ocr_text.strip():
                    confidence = min(0.85, 0.5 + len(ocr_text) / 5000)
                    results.append(ExtractionResult(
                        text=ocr_text, confidence=confidence,
                        method="vision_ocr",
                    ))
            except Exception as e:
                logger.debug(f"Vision OCR on PDF failed: {e}")

        return results

    # ═══════════════════════════════════════════
    #  HTML extraction — 2 methods
    # ═══════════════════════════════════════════

    def _extract_html(self, file_bytes):
        results = []
        encoding = self._detect_encoding(file_bytes)
        html_text = file_bytes.decode(encoding, errors="replace")

        # Method 1: BeautifulSoup — structured parsing
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "meta", "link"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            tables = self._extract_html_tables(soup)
            if text.strip():
                results.append(ExtractionResult(
                    text=text, confidence=0.85,
                    method="beautifulsoup", tables=tables,
                ))
        except Exception as e:
            logger.debug(f"BeautifulSoup failed: {e}")

        # Method 2: Regex strip — catches things BS4 sometimes misses
        try:
            raw = re.sub(r"<script[^>]*>.*?</script>", "", html_text, flags=re.DOTALL)
            raw = re.sub(r"<style[^>]*>.*?</style>", "", raw, flags=re.DOTALL)
            raw = re.sub(r"<[^>]+>", "\n", raw)
            raw = re.sub(r"\n{3,}", "\n\n", raw).strip()
            if raw and len(raw) > 20:
                results.append(ExtractionResult(
                    text=raw, confidence=0.6, method="raw_html_strip",
                ))
        except Exception as e:
            logger.debug(f"Raw HTML strip failed: {e}")

        return results

    # ═══════════════════════════════════════════
    #  Excel extraction — 2 methods
    # ═══════════════════════════════════════════

    def _extract_excel(self, file_bytes):
        results = []

        # Method 1: openpyxl — preserves formatting, cell-by-cell
        try:
            import openpyxl
            wb = openpyxl.load_workbook(
                io.BytesIO(file_bytes), read_only=True, data_only=True
            )
            all_text = []
            all_tables = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                all_text.append(f"[SHEET: {sheet_name}]")
                for row in ws.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_text):
                        rows.append(row_text)
                        all_text.append("\t".join(row_text))
                if rows:
                    all_tables.append(self._rows_to_dicts(rows))
            wb.close()
            full_text = "\n".join(all_text)
            if full_text.strip():
                results.append(ExtractionResult(
                    text=full_text, confidence=0.9,
                    method="openpyxl", tables=all_tables,
                ))
        except Exception as e:
            logger.debug(f"openpyxl failed: {e}")

        # Method 2: pandas (if available — handles merged cells, messy data)
        try:
            import pandas as pd
            dfs = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, dtype=str)
            all_text = []
            all_tables = []
            for sheet_name, df in dfs.items():
                df = df.fillna("")
                all_text.append(f"[SHEET: {sheet_name}]")
                all_text.append(df.to_string())
                all_tables.append(df.to_dict("records"))
            full_text = "\n".join(all_text)
            if full_text.strip():
                results.append(ExtractionResult(
                    text=full_text, confidence=0.85,
                    method="pandas_excel", tables=all_tables,
                ))
        except ImportError:
            pass  # pandas not installed — that's fine
        except Exception as e:
            logger.debug(f"pandas Excel failed: {e}")

        return results

    # ═══════════════════════════════════════════
    #  Word extraction
    # ═══════════════════════════════════════════

    def _extract_word(self, file_bytes):
        results = []

        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    paragraphs.append("\t".join(row_text))
            full_text = "\n".join(paragraphs)
            if full_text.strip():
                results.append(ExtractionResult(
                    text=full_text, confidence=0.9, method="python_docx",
                ))
        except Exception as e:
            logger.debug(f"python-docx failed: {e}")

        return results

    # ═══════════════════════════════════════════
    #  Image extraction
    # ═══════════════════════════════════════════

    def _extract_image(self, file_bytes):
        results = []

        # Method 1: Google Vision text_detection
        try:
            text = self._vision_ocr_image(file_bytes)
            if text and text.strip():
                results.append(ExtractionResult(
                    text=text, confidence=0.8, method="vision_ocr",
                ))
        except Exception as e:
            logger.debug(f"Vision OCR on image failed: {e}")

        # Method 2: Google Vision document_text_detection (better for dense docs)
        try:
            text = self._vision_document_ocr_image(file_bytes)
            if text and text.strip():
                results.append(ExtractionResult(
                    text=text, confidence=0.82, method="vision_document_ocr",
                ))
        except Exception as e:
            logger.debug(f"Vision document OCR failed: {e}")

        return results

    # ═══════════════════════════════════════════
    #  CSV extraction
    # ═══════════════════════════════════════════

    def _extract_csv(self, file_bytes):
        results = []
        encoding = self._detect_encoding(file_bytes)

        # Method 1: csv module
        try:
            text = file_bytes.decode(encoding, errors="replace")
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)
            full_text = "\n".join([",".join(row) for row in rows])
            tables = [self._rows_to_dicts(rows)] if len(rows) > 1 else []
            if full_text.strip():
                results.append(ExtractionResult(
                    text=full_text, confidence=0.9,
                    method="csv_module", tables=tables,
                ))
        except Exception as e:
            logger.debug(f"csv module failed: {e}")

        # Method 2: pandas (if available)
        try:
            import pandas as pd
            df = pd.read_csv(
                io.BytesIO(file_bytes), encoding=encoding,
                dtype=str, on_bad_lines="skip",
            )
            df = df.fillna("")
            full_text = df.to_string()
            tables = [df.to_dict("records")]
            if full_text.strip():
                results.append(ExtractionResult(
                    text=full_text, confidence=0.85,
                    method="pandas_csv", tables=tables,
                ))
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"pandas CSV failed: {e}")

        return results

    # ═══════════════════════════════════════════
    #  Fallbacks
    # ═══════════════════════════════════════════

    def _try_plain_text(self, file_bytes):
        """Try reading as plain text with encoding detection."""
        encoding = self._detect_encoding(file_bytes)
        try:
            text = file_bytes.decode(encoding, errors="replace")
            # Check if result looks like actual text vs. binary garbage
            garbage = sum(1 for c in text[:2000] if ord(c) < 32 and c not in "\n\r\t")
            sample_len = min(len(text), 2000)
            if sample_len > 0 and garbage / sample_len < 0.3 and text.strip():
                return ExtractionResult(
                    text=text, confidence=0.5, method="plain_text",
                )
        except Exception:
            pass
        return None

    def _try_vision_ocr_fallback(self, file_bytes, filename):
        """Last resort: send raw bytes to Vision OCR."""
        try:
            if filename.lower().endswith(".pdf"):
                text = self._vision_ocr_pdf(file_bytes)
            else:
                text = self._vision_ocr_image(file_bytes)
            if text and text.strip():
                return ExtractionResult(
                    text=text, confidence=0.5, method="vision_ocr_fallback",
                )
        except Exception:
            pass
        return None

    # ═══════════════════════════════════════════
    #  Google Vision OCR integration
    #  (uses the same pattern as rcb_helpers.py)
    # ═══════════════════════════════════════════

    def _vision_ocr_image(self, image_bytes):
        """OCR a single image using Google Vision text_detection."""
        from google.cloud import vision

        preprocessed = self._preprocess_image(image_bytes)
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=preprocessed)
        response = client.text_detection(image=image)

        if response.error.message:
            logger.warning(f"Vision API error: {response.error.message}")

        if response.text_annotations:
            return self._cleanup_hebrew(response.text_annotations[0].description)
        return ""

    def _vision_document_ocr_image(self, image_bytes):
        """OCR using document_text_detection — better for dense text/tables."""
        from google.cloud import vision

        preprocessed = self._preprocess_image(image_bytes)
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=preprocessed)
        response = client.document_text_detection(image=image)

        if response.error.message:
            logger.warning(f"Vision API error: {response.error.message}")

        if response.full_text_annotation:
            return self._cleanup_hebrew(response.full_text_annotation.text)
        return ""

    def _vision_ocr_pdf(self, pdf_bytes):
        """Convert PDF pages to images, preprocess, and OCR each page."""
        images = self._pdf_to_images(pdf_bytes)
        if not images:
            return ""

        all_text = []
        for i, img_bytes in enumerate(images):
            try:
                page_text = self._vision_ocr_image(img_bytes)
                if page_text:
                    if len(images) > 1:
                        all_text.append(f"--- Page {i + 1} ---\n{page_text}")
                    else:
                        all_text.append(page_text)
            except Exception as e:
                logger.debug(f"OCR page {i + 1} failed: {e}")

        return "\n\n".join(all_text)

    def _pdf_to_images(self, pdf_bytes, dpi=300, max_pages=20):
        """Convert PDF pages to PNG images using PyMuPDF (fitz)."""
        try:
            import fitz
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            images = []
            for page_num in range(min(len(doc), max_pages)):
                page = doc[page_num]
                mat = fitz.Matrix(dpi / 72, dpi / 72)
                pix = page.get_pixmap(matrix=mat)
                images.append(pix.tobytes("png"))
            doc.close()
            return images
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed — cannot convert PDF to images")
            return []
        except Exception as e:
            logger.debug(f"PDF to images failed: {e}")
            return []

    def _preprocess_image(self, img_bytes):
        """Preprocess image for better OCR accuracy (grayscale, contrast, sharpen)."""
        try:
            from PIL import Image, ImageEnhance, ImageFilter

            img = Image.open(io.BytesIO(img_bytes))
            if img.mode != "L":
                img = img.convert("L")
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.5)
            img = img.filter(ImageFilter.SHARPEN)

            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return buf.getvalue()
        except Exception:
            return img_bytes  # Return original if preprocessing fails

    # ═══════════════════════════════════════════
    #  AI-assisted reading (expensive — last resort)
    # ═══════════════════════════════════════════

    def try_ai_document_reading(self, file_bytes, filename, content_type):
        """
        Send document to Gemini Vision and ask it to read everything.
        EXPENSIVE — only use when all other methods fail or return low confidence.

        Returns ExtractionResult or None.
        """
        try:
            import base64

            # For images or small PDFs, convert to base64 for Gemini
            if filename.lower().endswith(".pdf"):
                images = self._pdf_to_images(file_bytes, dpi=200, max_pages=5)
                if not images:
                    return None
                # Use first page image for now
                b64 = base64.b64encode(images[0]).decode()
            else:
                b64 = base64.b64encode(file_bytes).decode()

            # Placeholder: wire to actual Gemini/Claude vision API
            # The integration session will connect this to the existing AI callers.
            #
            # prompt = (
            #     "Read this document completely. Extract ALL text content including "
            #     "paragraphs, tables, numbers, codes, Hebrew and English text. "
            #     "For tables, format as: HEADER1 | HEADER2 | HEADER3\\n"
            #     "value1 | value2 | value3. Do NOT summarize — extract EVERYTHING verbatim."
            # )
            # result = call_gemini_vision(b64, prompt)
            # if result:
            #     return ExtractionResult(
            #         text=result, confidence=0.75, method="gemini_vision_read"
            #     )
            return None
        except Exception as e:
            logger.debug(f"AI document reading failed: {e}")
            return None

    # ═══════════════════════════════════════════
    #  Helper functions
    # ═══════════════════════════════════════════

    def _pick_best(self, results):
        """Compare results from multiple methods, pick best, flag disagreements."""
        results.sort(key=lambda r: r.confidence, reverse=True)
        best = results[0]

        if len(results) >= 2:
            text_1 = results[0].text.strip()
            text_2 = results[1].text.strip()

            similarity = self._text_similarity(text_1, text_2)

            if similarity < 0.5 and len(text_1) > 100 and len(text_2) > 100:
                best.warnings.append(
                    f"Methods disagree ({similarity:.0%} similarity). "
                    f"'{results[0].method}': {len(text_1)} chars, "
                    f"'{results[1].method}': {len(text_2)} chars. "
                    f"Using '{results[0].method}'. Review recommended."
                )
                best.needs_review = True

            # If second method found significantly more text, prefer it
            if (
                len(text_2) > len(text_1) * 1.5
                and results[1].confidence > results[0].confidence * 0.8
            ):
                best = results[1]
                best.warnings.append(
                    f"Switched to '{results[1].method}' — found more content "
                    f"({len(text_2)} vs {len(text_1)} chars)."
                )

        best.methods_tried = [r.method for r in results]
        best.all_results_summary = [
            {"method": r.method, "confidence": round(r.confidence, 2), "chars": len(r.text)}
            for r in results
        ]
        return best

    def _text_similarity(self, text1, text2):
        """Quick similarity check using word overlap."""
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        if not words1 or not words2:
            return 0.0
        overlap = words1 & words2
        return len(overlap) / max(len(words1), len(words2))

    def _detect_encoding(self, file_bytes):
        """
        Detect file encoding. Important for Hebrew content.
        Israeli government files often use Windows-1255, not UTF-8.
        """
        # Try chardet if available
        try:
            import chardet
            result = chardet.detect(file_bytes[:10000])
            enc = result.get("encoding")
            if enc:
                return enc
        except ImportError:
            pass

        # Manual detection: try common encodings for Hebrew
        for enc in ("utf-8", "windows-1255", "iso-8859-8", "utf-16", "latin-1"):
            try:
                file_bytes[:2000].decode(enc)
                return enc
            except (UnicodeDecodeError, LookupError):
                continue
        return "utf-8"

    def _cleanup_hebrew(self, text):
        """Fix common Hebrew/RTL extraction issues."""
        if not text:
            return text
        text = re.sub(r" {3,}", "  ", text)
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        replacements = {
            "חשבוו": "חשבון",
            'מע"ט': 'מע"מ',
            "עמיל מכם": "עמיל מכס",
            "מתווך מכס": "עמיל מכס",
            "מתווכי מכס": "עמילי מכס",
        }
        for wrong, right in replacements.items():
            text = text.replace(wrong, right)
        return text

    def _table_to_dicts(self, table_rows):
        """Convert raw table rows to list of dicts using first row as headers."""
        if not table_rows or len(table_rows) < 2:
            return []
        headers = [
            str(h).strip() if h else f"col_{i}"
            for i, h in enumerate(table_rows[0])
        ]
        return [
            {
                headers[i] if i < len(headers) else f"col_{i}": str(cell).strip() if cell else ""
                for i, cell in enumerate(row)
            }
            for row in table_rows[1:]
            if any(cell for cell in row)
        ]

    def _rows_to_dicts(self, rows):
        """Same as _table_to_dicts but handles edge cases."""
        return self._table_to_dicts(rows)

    def _extract_html_tables(self, soup):
        """Extract all <table> elements as structured data."""
        tables = []
        for table_el in soup.find_all("table"):
            rows = []
            for tr in table_el.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            if rows:
                parsed = self._table_to_dicts(rows)
                if parsed:
                    tables.append(parsed)
        return tables
