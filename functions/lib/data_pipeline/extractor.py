"""
Step 2: EXTRACT — Convert raw files to plain text.

Every format → plain text + tables. The brain reads TEXT, not PDFs.

Supported formats:
- PDF (pdfplumber → PyMuPDF fallback → Cloud Vision OCR fallback)
- HTML (BeautifulSoup → clean text + table extraction)
- Word (.docx via python-docx)
- Excel (.xlsx via openpyxl)
- Images (Cloud Vision OCR)
- Plain text (direct)

Session 27 — Assignment 14C
"""

import logging
import re

logger = logging.getLogger("rcb.pipeline.extractor")


def extract_text(file_bytes, content_type, filename=""):
    """
    Extract text from raw file bytes.

    Args:
        file_bytes: bytes — raw file content
        content_type: str — MIME type or file extension
        filename: str — original filename (for format detection fallback)

    Returns:
        dict with:
            full_text: str — extracted plain text
            tables: list — extracted tables (list of list-of-dicts)
            language: str — detected primary language ("he", "en", "mixed")
            extraction_method: str — which extractor was used
            char_count: int — length of extracted text
    """
    if not file_bytes:
        return _empty_result("empty_input")

    # Normalize content_type
    ct = (content_type or "").lower()
    fn = (filename or "").lower()

    # Route to appropriate extractor
    if "pdf" in ct or fn.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif "html" in ct or fn.endswith((".html", ".htm")):
        return _extract_html(file_bytes)
    elif "wordprocessingml" in ct or "msword" in ct or fn.endswith((".docx", ".doc")):
        return _extract_docx(file_bytes)
    elif "spreadsheetml" in ct or "excel" in ct or fn.endswith((".xlsx", ".xls")):
        return _extract_excel(file_bytes)
    elif ct.startswith("image/") or fn.endswith((".png", ".jpg", ".jpeg", ".tiff")):
        return _extract_image_ocr(file_bytes)
    elif "text" in ct or "json" in ct or fn.endswith((".txt", ".csv", ".json")):
        return _extract_plain_text(file_bytes)
    else:
        # Try as plain text
        return _extract_plain_text(file_bytes)


# ═══════════════════════════════════════════
#  PDF EXTRACTION (3-level fallback)
# ═══════════════════════════════════════════

def _extract_pdf(file_bytes):
    """Extract text from PDF with 3-level fallback."""
    # Level 1: pdfplumber (best for text-based PDFs)
    try:
        import pdfplumber
        import io

        text_parts = []
        tables = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages[:50]:  # Max 50 pages
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                # Extract tables
                page_tables = page.extract_tables()
                for tbl in (page_tables or []):
                    if tbl and len(tbl) > 1:
                        headers = [str(h or "").strip() for h in tbl[0]]
                        for row in tbl[1:]:
                            row_dict = {}
                            for i, cell in enumerate(row):
                                key = headers[i] if i < len(headers) else f"col_{i}"
                                row_dict[key] = str(cell or "").strip()
                            tables.append(row_dict)

        full_text = "\n".join(text_parts)
        if full_text and len(full_text.strip()) > 50:
            return _build_result(full_text, tables, "pdfplumber")
    except Exception as e:
        logger.debug(f"pdfplumber failed: {e}")

    # Level 2: PyMuPDF (better for some PDF types)
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc[:50]:
            text_parts.append(page.get_text())
        doc.close()

        full_text = "\n".join(text_parts)
        if full_text and len(full_text.strip()) > 50:
            return _build_result(full_text, [], "pymupdf")
    except Exception as e:
        logger.debug(f"PyMuPDF failed: {e}")

    # Level 3: OCR fallback (for scanned PDFs)
    return _extract_image_ocr(file_bytes, source="pdf_ocr")


# ═══════════════════════════════════════════
#  HTML EXTRACTION
# ═══════════════════════════════════════════

def _extract_html(file_bytes):
    """Extract text and tables from HTML."""
    try:
        html = file_bytes.decode("utf-8", errors="replace")
    except Exception:
        html = str(file_bytes)

    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
    except ImportError:
        # Fallback: regex-based HTML stripping
        text = re.sub(r'<[^>]+>', ' ', html)
        text = re.sub(r'\s+', ' ', text).strip()
        return _build_result(text, [], "regex_strip")

    # Remove non-content elements
    for tag in soup(["script", "style", "nav", "footer", "header", "meta", "link"]):
        tag.decompose()

    # Extract tables
    tables = []
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if len(rows) > 1:
            headers = [th.get_text(strip=True) for th in rows[0].find_all(["th", "td"])]
            for row in rows[1:]:
                cells = [td.get_text(strip=True) for td in row.find_all(["td", "th"])]
                row_dict = {}
                for i, cell in enumerate(cells):
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    row_dict[key] = cell
                if any(row_dict.values()):
                    tables.append(row_dict)

    # Extract text
    text = soup.get_text(separator="\n", strip=True)

    return _build_result(text, tables, "beautifulsoup")


# ═══════════════════════════════════════════
#  WORD DOCUMENT EXTRACTION
# ═══════════════════════════════════════════

def _extract_docx(file_bytes):
    """Extract text from .docx files."""
    try:
        import docx
        import io

        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        tables = []
        for table in doc.tables:
            if len(table.rows) > 1:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                for row in table.rows[1:]:
                    row_dict = {}
                    for i, cell in enumerate(row.cells):
                        key = headers[i] if i < len(headers) else f"col_{i}"
                        row_dict[key] = cell.text.strip()
                    if any(row_dict.values()):
                        tables.append(row_dict)

        full_text = "\n".join(text_parts)
        return _build_result(full_text, tables, "python_docx")
    except Exception as e:
        logger.warning(f"docx extraction failed: {e}")
        return _empty_result("docx_failed")


# ═══════════════════════════════════════════
#  EXCEL EXTRACTION
# ═══════════════════════════════════════════

def _extract_excel(file_bytes):
    """Extract text from Excel files."""
    try:
        import openpyxl
        import io

        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)

        text_parts = []
        tables = []

        for sheet_name in wb.sheetnames[:5]:  # Max 5 sheets
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(max_row=500, values_only=True))

            if not rows:
                continue

            headers = [str(h or "").strip() for h in rows[0]]
            text_parts.append(f"Sheet: {sheet_name}")
            text_parts.append(" | ".join(headers))

            for row in rows[1:]:
                row_dict = {}
                row_text_parts = []
                for i, cell in enumerate(row):
                    val = str(cell or "").strip()
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    row_dict[key] = val
                    if val:
                        row_text_parts.append(val)
                if any(row_dict.values()):
                    tables.append(row_dict)
                if row_text_parts:
                    text_parts.append(" | ".join(row_text_parts))

        wb.close()
        full_text = "\n".join(text_parts)
        return _build_result(full_text, tables, "openpyxl")
    except Exception as e:
        logger.warning(f"Excel extraction failed: {e}")
        return _empty_result("excel_failed")


# ═══════════════════════════════════════════
#  IMAGE OCR
# ═══════════════════════════════════════════

def _extract_image_ocr(file_bytes, source="image_ocr"):
    """Extract text from images using Google Cloud Vision OCR."""
    try:
        from google.cloud import vision

        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=file_bytes)
        response = client.text_detection(image=image)

        if response.text_annotations:
            full_text = response.text_annotations[0].description
            return _build_result(full_text, [], f"vision_{source}")
    except ImportError:
        logger.debug("google-cloud-vision not available for OCR")
    except Exception as e:
        logger.warning(f"OCR failed: {e}")

    return _empty_result(f"{source}_failed")


# ═══════════════════════════════════════════
#  PLAIN TEXT
# ═══════════════════════════════════════════

def _extract_plain_text(file_bytes):
    """Handle plain text, CSV, JSON files."""
    try:
        text = file_bytes.decode("utf-8", errors="replace")
    except Exception:
        text = str(file_bytes)[:50000]

    return _build_result(text, [], "plain_text")


# ═══════════════════════════════════════════
#  HELPERS
# ═══════════════════════════════════════════

def _build_result(full_text, tables, method):
    """Build a standard extraction result."""
    text = (full_text or "").strip()
    capped = text[:200000]  # Cap at 200K chars
    language = _detect_language(capped)
    return {
        "full_text": capped,
        "tables": tables[:500],  # Cap at 500 table rows
        "language": language,
        "extraction_method": method,
        "char_count": len(capped),
    }


def _empty_result(method):
    """Return empty extraction result."""
    return {
        "full_text": "",
        "tables": [],
        "language": "unknown",
        "extraction_method": method,
        "char_count": 0,
    }


def _detect_language(text):
    """Simple language detection: Hebrew vs English vs mixed."""
    if not text:
        return "unknown"

    sample = text[:2000]
    hebrew_chars = sum(1 for c in sample if "\u0590" <= c <= "\u05FF")
    latin_chars = sum(1 for c in sample if "a" <= c.lower() <= "z")
    total = hebrew_chars + latin_chars

    if total == 0:
        return "unknown"

    hebrew_ratio = hebrew_chars / total

    if hebrew_ratio > 0.7:
        return "he"
    elif hebrew_ratio < 0.3:
        return "en"
    else:
        return "mixed"
