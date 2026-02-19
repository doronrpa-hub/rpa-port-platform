"""RCB Helper functions - Graph API, PDF extraction, Hebrew names
UPDATED Session 17 Phase 0: Improved document extraction
"""
import requests
import base64
import io
import re
import hashlib
from datetime import datetime, timedelta, timezone

# ============================================================
# PDF TEXT EXTRACTION (with OCR fallback)
# Phase 0: Improved extraction quality, table handling, OCR
# ============================================================

def _assess_extraction_quality(text):
    """Check if extracted text is actually useful, not garbage"""
    if not text or len(text.strip()) < 50:
        return False, "too_short"

    stripped = text.strip()

    # Check for Hebrew content (customs docs should have Hebrew)
    hebrew_chars = sum(1 for c in stripped if '\u0590' <= c <= '\u05FF')
    latin_chars = sum(1 for c in stripped if c.isascii() and c.isalpha())
    digit_chars = sum(1 for c in stripped if c.isdigit())
    total_meaningful = hebrew_chars + latin_chars + digit_chars

    # If less than 30% of text is meaningful characters, it's garbage
    if len(stripped) > 0 and total_meaningful / len(stripped) < 0.3:
        return False, "garbage_chars"

    # Check for invoice-like patterns (numbers, currency, dates)
    has_numbers = bool(re.search(r'\d{3,}', stripped))  # 3+ digit numbers
    has_currency = bool(re.search(r'(USD|EUR|ILS|NIS|\$|€|₪)', stripped, re.IGNORECASE))
    has_date = bool(re.search(r'\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}', stripped))

    # If we have meaningful length but no numbers AND no meaningful text, probably wrong extraction
    # Don't reject certificates, EUR.1, letters that have text but no numbers
    if len(stripped) > 200 and not has_numbers and hebrew_chars < 20 and latin_chars < 20:
        return False, "no_numbers"

    return True, "ok"


def _cleanup_hebrew_text(text):
    """Fix common Hebrew/RTL extraction issues"""
    if not text:
        return text

    # Normalize whitespace
    text = re.sub(r' {3,}', '  ', text)  # Multiple spaces to double
    text = re.sub(r'\n{4,}', '\n\n\n', text)  # Multiple newlines to max 3

    # Fix common OCR mistakes in Hebrew customs context
    replacements = {
        'חשבוו': 'חשבון',
        'מע"ט': 'מע"מ',
        'עמיל מכם': 'עמיל מכס',
    }
    for wrong, right in replacements.items():
        text = text.replace(wrong, right)

    return text


def _tag_document_structure(text):
    """Tag document sections to help AI classification agent"""
    tags = []

    # Detect invoice number
    inv_match = re.search(r'(?:invoice|inv|חשבונית|חשבון)[\s#:]*(\S+)', text, re.IGNORECASE)
    if inv_match:
        tags.append(f"[INVOICE_NUMBER: {inv_match.group(1)}]")

    # Detect dates
    dates = re.findall(r'\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}', text)
    if dates:
        tags.append(f"[DATES_FOUND: {', '.join(dates[:5])}]")

    # Detect currency and amounts
    amounts = re.findall(r'(?:USD|EUR|ILS|NIS|\$|€|₪)\s*[\d,.]+', text, re.IGNORECASE)
    if not amounts:
        amounts = re.findall(r'[\d,.]+\s*(?:USD|EUR|ILS|NIS|\$|€|₪)', text, re.IGNORECASE)
    if amounts:
        tags.append(f"[AMOUNTS: {', '.join(amounts[:10])}]")

    # Detect HS code patterns (XX.XX or XXXX.XX.XXXX)
    hs_codes = re.findall(r'\b\d{4}[.\s]\d{2}(?:[.\s]\d{2,6})?\b', text)
    if hs_codes:
        tags.append(f"[HS_CODE_CANDIDATES: {', '.join(set(hs_codes[:10]))}]")

    # Detect BL/AWB numbers
    bl_match = re.search(r'(?:B/?L|bill\s*of\s*lading|שטר\s*מטען)[\s#:]*(\S+)', text, re.IGNORECASE)
    if bl_match:
        tags.append(f"[BL_NUMBER: {bl_match.group(1)}]")

    awb_match = re.search(r'(?:AWB|air\s*waybill)[\s#:]*(\S+)', text, re.IGNORECASE)
    if awb_match:
        tags.append(f"[AWB_NUMBER: {awb_match.group(1)}]")

    # Detect country names
    countries = re.findall(r'(?:China|India|Turkey|Germany|USA|Italy|Japan|Korea|Vietnam|Thailand|'
                           r'סין|הודו|טורקיה|גרמניה|ארה"ב|איטליה|יפן|קוריאה|וייטנאם|תאילנד)',
                           text, re.IGNORECASE)
    if countries:
        tags.append(f"[COUNTRIES: {', '.join(set(countries[:5]))}]")

    # Detect shipping terms
    incoterms = re.findall(r'\b(?:FOB|CIF|CFR|EXW|DDP|DAP|FCA|CPT|CIP|DAT)\b', text, re.IGNORECASE)
    if incoterms:
        tags.append(f"[INCOTERMS: {', '.join(set(incoterms))}]")

    if tags:
        header = "\n".join(tags)
        return f"[DOCUMENT_ANALYSIS]\n{header}\n[/DOCUMENT_ANALYSIS]\n\n{text}"

    return text


def _preprocess_image_for_ocr(img_bytes):
    """Preprocess image for better OCR accuracy"""
    try:
        from PIL import Image, ImageEnhance, ImageFilter

        img = Image.open(io.BytesIO(img_bytes))

        # Convert to grayscale
        if img.mode != 'L':
            img = img.convert('L')

        # Enhance contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)

        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)

        # Convert back to bytes
        output = io.BytesIO()
        img.save(output, format='PNG')
        return output.getvalue()
    except Exception as e:
        print(f"  Image preprocessing error: {e}")
        return img_bytes  # Return original if preprocessing fails


def extract_text_from_pdf_bytes(pdf_bytes):
    """Extract text from PDF bytes - tries multiple methods
    Phase 0: Improved quality assessment, combined extraction, better OCR"""

    # Method 1: pdfplumber (fast, for text-based PDFs)
    text = _extract_with_pdfplumber(pdf_bytes)
    is_good, reason = _assess_extraction_quality(text)
    if is_good:
        print(f"    ✅ pdfplumber extracted {len(text)} chars (quality: {reason})")
        return text
    elif text:
        print(f"    ⚠️ pdfplumber: {len(text)} chars but quality: {reason}")

    # Method 2: pypdf (sometimes works better)
    text = _extract_with_pypdf(pdf_bytes)
    is_good, reason = _assess_extraction_quality(text)
    if is_good:
        print(f"    ✅ pypdf extracted {len(text)} chars (quality: {reason})")
        return text

    # Method 3: COMBINED - try pdfplumber + pypdf merged
    combined = _extract_with_pdfplumber(pdf_bytes) + "\n" + _extract_with_pypdf(pdf_bytes)
    is_good, reason = _assess_extraction_quality(combined)
    if is_good:
        print(f"    ✅ Combined extraction: {len(combined)} chars")
        return combined

    # Method 4: Vision OCR (preprocessed, 300 DPI)
    print(f"    🔍 All text methods failed, trying OCR...")
    text = _extract_with_vision_ocr(pdf_bytes)
    if text:
        print(f"    ✅ OCR extracted {len(text)} chars")
    else:
        print(f"    ⚠️ All extraction methods failed")
    return text or ""


def _extract_with_pdfplumber(pdf_bytes):
    """Extract text using pdfplumber with improved table handling"""
    try:
        import pdfplumber
        text_parts = []

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract regular text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Extract tables with better settings
                tables = page.extract_tables({
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 5,
                    "join_tolerance": 5,
                    "edge_min_length": 10,
                    "min_words_vertical": 1,
                    "min_words_horizontal": 1,
                })

                for table_idx, table in enumerate(tables):
                    if not table:
                        continue

                    # Format table with structure preserved
                    table_text = f"\n[TABLE {table_idx + 1} on page {page_num + 1}]\n"
                    for row in table:
                        if row:
                            cells = [str(cell).strip() if cell else "" for cell in row]
                            # Skip completely empty rows
                            if any(cells):
                                table_text += " | ".join(cells) + "\n"
                    table_text += f"[/TABLE]\n"
                    text_parts.append(table_text)

        return _cleanup_hebrew_text("\n".join(text_parts))
    except Exception as e:
        print(f"    pdfplumber error: {e}")
        return ""


def _extract_with_pypdf(pdf_bytes):
    """Extract text using pypdf"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return _cleanup_hebrew_text("\n".join(text_parts))
    except Exception as e:
        print(f"    pypdf error: {e}")
        return ""


def _extract_with_vision_ocr(pdf_bytes):
    """Extract text using Google Cloud Vision OCR
    Phase 0: Added image preprocessing before OCR"""
    try:
        from google.cloud import vision

        # Convert PDF pages to images first
        images = _pdf_to_images(pdf_bytes)
        if not images:
            print(f"    ⚠️ Could not convert PDF to images")
            return ""

        client = vision.ImageAnnotatorClient()
        all_text = []

        for i, img_bytes in enumerate(images):
            # Phase 0: Preprocess image before OCR
            img_bytes = _preprocess_image_for_ocr(img_bytes)
            image = vision.Image(content=img_bytes)
            response = client.text_detection(image=image)

            if response.text_annotations:
                page_text = response.text_annotations[0].description
                all_text.append(f"--- Page {i+1} ---\n{page_text}")
                print(f"    📄 Page {i+1}: {len(page_text)} chars")

            if response.error.message:
                print(f"    Vision API error: {response.error.message}")

        return _cleanup_hebrew_text("\n\n".join(all_text))
    except ImportError:
        print(f"    ⚠️ google-cloud-vision not installed")
        return ""
    except Exception as e:
        print(f"    Vision OCR error: {e}")
        return ""


def _pdf_to_images(pdf_bytes):
    """Convert PDF pages to images for OCR
    Phase 0: Raised DPI from 150 to 300"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        images = []

        for page_num in range(min(len(doc), 10)):  # Max 10 pages
            page = doc[page_num]
            # Phase 0: Render at 300 DPI (was 150) for better OCR accuracy
            pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
            img_bytes = pix.tobytes("png")
            images.append(img_bytes)

        doc.close()
        return images
    except ImportError:
        print(f"    ⚠️ PyMuPDF not installed, trying pdf2image...")
        return _pdf_to_images_fallback(pdf_bytes)
    except Exception as e:
        print(f"    PDF to image error: {e}")
        return []


def _pdf_to_images_fallback(pdf_bytes):
    """Fallback: Convert PDF to images using pdf2image"""
    try:
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(pdf_bytes, dpi=300, first_page=1, last_page=10)
        result = []
        for img in images:
            buf = io.BytesIO()
            img.save(buf, format='PNG')
            result.append(buf.getvalue())
        return result
    except Exception as e:
        print(f"    pdf2image error: {e}")
        return []


def _try_decode(raw_bytes):
    """Decode bytes trying Hebrew-friendly encodings.
    Israeli government files often use Windows-1255, not UTF-8."""
    for encoding in ('utf-8', 'windows-1255', 'iso-8859-8', 'latin-1'):
        try:
            return raw_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw_bytes.decode('utf-8', errors='replace')


def _extract_from_excel(file_bytes):
    """Extract text from Excel files (.xlsx, .xls)"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"[SHEET: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(cells):
                    text_parts.append(" | ".join(cells))
            text_parts.append(f"[/SHEET]")

        wb.close()
        return "\n".join(text_parts)
    except Exception as e:
        print(f"    Excel extraction error: {e}")
        return ""


def _extract_from_docx(file_bytes):
    """Extract text from Word documents (.docx)"""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            text_parts.append(f"\n[TABLE {table_idx + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    text_parts.append(" | ".join(cells))
            text_parts.append(f"[/TABLE]")

        return "\n".join(text_parts)
    except Exception as e:
        print(f"    Word extraction error: {e}")
        return ""


def _extract_from_eml(file_bytes):
    """Extract text from email files (.eml, .msg)"""
    try:
        import email
        from email import policy
        msg = email.message_from_bytes(file_bytes, policy=policy.default)
        text_parts = []

        # Extract headers
        subject = msg.get('Subject', '')
        from_addr = msg.get('From', '')
        date = msg.get('Date', '')
        if subject:
            text_parts.append(f"Subject: {subject}")
        if from_addr:
            text_parts.append(f"From: {from_addr}")
        if date:
            text_parts.append(f"Date: {date}")
        text_parts.append("")

        # Extract body
        body = msg.get_body(preferencelist=('plain', 'html'))
        if body:
            content = body.get_content()
            if body.get_content_type() == 'text/html':
                # Strip HTML tags
                content = re.sub(r'<[^>]+>', ' ', content)
                content = re.sub(r'\s+', ' ', content).strip()
            text_parts.append(content)

        return "\n".join(text_parts)
    except Exception as e:
        print(f"    Email extraction error: {e}")
        return ""


def _extract_from_msg(file_bytes):
    """Extract text from Outlook .msg files"""
    try:
        import extract_msg
        msg = extract_msg.Message(io.BytesIO(file_bytes))
        text_parts = []

        if msg.subject:
            text_parts.append(f"Subject: {msg.subject}")
        if msg.sender:
            text_parts.append(f"From: {msg.sender}")
        if msg.date:
            text_parts.append(f"Date: {msg.date}")
        text_parts.append("")

        if msg.body:
            text_parts.append(msg.body)
        elif msg.htmlBody:
            content = re.sub(r'<[^>]+>', ' ', msg.htmlBody if isinstance(msg.htmlBody, str) else _try_decode(msg.htmlBody))
            content = re.sub(r'\s+', ' ', content).strip()
            text_parts.append(content)

        msg.close()
        return "\n".join(text_parts)
    except ImportError:
        print("    extract-msg not installed, cannot read .msg files")
        return ""
    except Exception as e:
        print(f"    MSG extraction error: {e}")
        return ""


def _extract_urls_from_text(text):
    """Detect URLs in text and append them as tagged section"""
    if not text:
        return text
    urls = re.findall(r'https?://[^\s<>"\')\]]+', text)
    if urls:
        unique_urls = list(dict.fromkeys(urls))  # Preserve order, remove dupes
        url_section = "\n[URLS_FOUND]\n" + "\n".join(unique_urls[:20]) + "\n[/URLS_FOUND]"
        return text + url_section
    return text


def extract_text_from_attachments(attachments_data, email_body=None):
    """Extract text from all attachments (PDF, Excel, Word, images, emails)
    Phase 0: Added Hebrew cleanup, structure tagging, multi-format support"""
    all_text = []

    # Extract URLs from email body if provided
    if email_body:
        email_body = _extract_urls_from_text(email_body)
        all_text.append(f"=== Email Body ===\n{email_body}")

    for att in attachments_data:
        name = att.get('name', 'file')
        content_bytes = att.get('contentBytes', '')

        if not content_bytes:
            continue

        # Decode base64
        try:
            file_bytes = base64.b64decode(content_bytes)
        except Exception:
            continue

        name_lower = name.lower()

        # PDF
        if name_lower.endswith('.pdf'):
            print(f"    📄 Extracting text from: {name}")
            text = extract_text_from_pdf_bytes(file_bytes)
            if text:
                text = _cleanup_hebrew_text(text)
                text = _tag_document_structure(text)
                all_text.append(f"=== {name} ===\n{text}")
            else:
                all_text.append(f"=== {name} ===\n[No text could be extracted]")

        # Excel
        elif name_lower.endswith(('.xlsx', '.xls')):
            print(f"    📊 Extracting text from: {name}")
            text = _extract_from_excel(file_bytes)
            if text:
                text = _cleanup_hebrew_text(text)
                text = _tag_document_structure(text)
                all_text.append(f"=== {name} ===\n{text}")

        # Word
        elif name_lower.endswith('.docx'):
            print(f"    📝 Extracting text from: {name}")
            text = _extract_from_docx(file_bytes)
            if text:
                text = _cleanup_hebrew_text(text)
                text = _tag_document_structure(text)
                all_text.append(f"=== {name} ===\n{text}")

        # Email files — .eml
        elif name_lower.endswith('.eml'):
            print(f"    📧 Extracting text from: {name}")
            text = _extract_from_eml(file_bytes)
            if text:
                text = _cleanup_hebrew_text(text)
                text = _extract_urls_from_text(text)
                all_text.append(f"=== {name} ===\n{text}")

        # Email files — Outlook .msg
        elif name_lower.endswith('.msg'):
            print(f"    📧 Extracting text from: {name}")
            text = _extract_from_msg(file_bytes)
            if text:
                text = _cleanup_hebrew_text(text)
                text = _extract_urls_from_text(text)
                all_text.append(f"=== {name} ===\n{text}")

        # Images — OCR
        elif name_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
            print(f"    🖼️ OCR on image: {name}")
            img_bytes = _preprocess_image_for_ocr(file_bytes)
            text = _ocr_image(img_bytes)
            if text:
                text = _cleanup_hebrew_text(text)
                text = _tag_document_structure(text)
                all_text.append(f"=== {name} ===\n{text}")

        # TIFF — multi-page image, convert each page then OCR
        elif name_lower.endswith(('.tiff', '.tif')):
            print(f"    🖼️ OCR on TIFF: {name}")
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(file_bytes))
                tiff_parts = []
                page = 0
                while True:
                    try:
                        img.seek(page)
                    except EOFError:
                        break
                    buf = io.BytesIO()
                    img.save(buf, format='PNG')
                    page_bytes = _preprocess_image_for_ocr(buf.getvalue())
                    page_text = _ocr_image(page_bytes)
                    if page_text:
                        tiff_parts.append(f"--- Page {page + 1} ---\n{page_text}")
                    page += 1
                if tiff_parts:
                    text = _cleanup_hebrew_text("\n\n".join(tiff_parts))
                    text = _tag_document_structure(text)
                    all_text.append(f"=== {name} ===\n{text}")
            except Exception as e:
                print(f"    TIFF extraction error: {e}")

        # CSV / TSV
        elif name_lower.endswith(('.csv', '.tsv')):
            print(f"    📊 Extracting text from: {name}")
            try:
                text = _try_decode(file_bytes)
                sep = '\t' if name_lower.endswith('.tsv') else ','
                lines = text.strip().split('\n')
                table_parts = [f"[TABLE CSV]"]
                for line in lines[:500]:  # Cap at 500 rows
                    cells = line.strip().split(sep)
                    if any(c.strip() for c in cells):
                        table_parts.append(" | ".join(c.strip() for c in cells))
                table_parts.append("[/TABLE]")
                result_text = "\n".join(table_parts)
                result_text = _cleanup_hebrew_text(result_text)
                result_text = _tag_document_structure(result_text)
                all_text.append(f"=== {name} ===\n{result_text}")
            except Exception as e:
                print(f"    CSV extraction error: {e}")

        # HTML
        elif name_lower.endswith(('.html', '.htm')):
            print(f"    🌐 Extracting text from: {name}")
            try:
                text = _try_decode(file_bytes)
                text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
                text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
                text = re.sub(r'<[^>]+>', ' ', text)
                text = re.sub(r'\s+', ' ', text).strip()
                if text:
                    text = _cleanup_hebrew_text(text)
                    text = _tag_document_structure(text)
                    all_text.append(f"=== {name} ===\n{text}")
            except Exception as e:
                print(f"    HTML extraction error: {e}")

    # Fix 6: Extraction summary logging
    print(f"  📊 Extraction summary: {len(all_text)} files, {sum(len(t) for t in all_text)} total chars")
    for entry in all_text:
        # Extract filename from "=== filename ===" header
        header_end = entry.find('\n')
        header = entry[:header_end] if header_end > 0 else entry[:50]
        content = entry[header_end+1:] if header_end > 0 else ""
        preview = content[:100].replace('\n', ' ').strip()
        print(f"    {header} ({len(content)} chars) -> {preview}...")

    return "\n\n".join(all_text)


def _ocr_image(image_bytes):
    """OCR a single image using Google Cloud Vision"""
    try:
        from google.cloud import vision
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=image_bytes)
        response = client.text_detection(image=image)
        
        if response.text_annotations:
            return response.text_annotations[0].description
        return ""
    except Exception as e:
        print(f"    Image OCR error: {e}")
        return ""


# ============================================================
# MICROSOFT GRAPH API HELPERS
# ============================================================

def get_rcb_secrets_internal(get_secret_func):
    """Get RCB email secrets from Secret Manager"""
    try:
        secrets = {}
        for name in ['RCB_EMAIL', 'RCB_PASSWORD', 'RCB_FALLBACK_EMAIL', 
                     'RCB_GRAPH_CLIENT_ID', 'RCB_GRAPH_TENANT_ID', 'RCB_GRAPH_CLIENT_SECRET',
                     'ANTHROPIC_API_KEY']:
            secrets[name] = get_secret_func(name)
        return secrets if secrets.get('RCB_EMAIL') else None
    except Exception as e:
        print(f"Secret error: {e}")
        return None

def helper_get_graph_token(secrets):
    """Get Microsoft Graph API access token"""
    try:
        response = requests.post(
            f"https://login.microsoftonline.com/{secrets.get('RCB_GRAPH_TENANT_ID')}/oauth2/v2.0/token",
            data={
                'client_id': secrets.get('RCB_GRAPH_CLIENT_ID'),
                'client_secret': secrets.get('RCB_GRAPH_CLIENT_SECRET'),
                'scope': 'https://graph.microsoft.com/.default',
                'grant_type': 'client_credentials'
            }
        )
        return response.json().get('access_token') if response.status_code == 200 else None
    except Exception as e:
        print(f"Token error: {e}")
        return None

def helper_graph_messages(access_token, user_email, unread_only=True, max_results=10):
    """Get messages from inbox"""
    try:
        url = f"https://graph.microsoft.com/v1.0/users/{user_email}/mailFolders/inbox/messages"
        params = {'$top': max_results, '$orderby': 'receivedDateTime desc'}
        if unread_only:
            params['$filter'] = 'isRead eq false'
        response = requests.get(url, headers={'Authorization': f'Bearer {access_token}'}, params=params)
        return response.json().get('value', []) if response.status_code == 200 else []
    except Exception as e:
        print(f"Graph messages error: {e}")
        return []

def helper_graph_attachments(access_token, user_email, message_id):
    """Get message attachments"""
    try:
        url = f"https://graph.microsoft.com/v1.0/users/{user_email}/messages/{message_id}/attachments"
        response = requests.get(url, headers={'Authorization': f'Bearer {access_token}'})
        return response.json().get('value', []) if response.status_code == 200 else []
    except Exception as e:
        print(f"Graph attachments error: {e}")
        return []

def helper_graph_mark_read(access_token, user_email, message_id):
    """Mark message as read"""
    try:
        url = f"https://graph.microsoft.com/v1.0/users/{user_email}/messages/{message_id}"
        requests.patch(url, headers={'Authorization': f'Bearer {access_token}', 'Content-Type': 'application/json'}, json={'isRead': True})
    except Exception as e:
        print(f"Graph mark-read error: {e}")

def _is_internal_recipient(email: str) -> bool:
    """Only @rpa-port.co.il recipients are allowed."""
    return (email or "").strip().lower().endswith("@rpa-port.co.il")


def helper_graph_send(access_token, user_email, to_email, subject, body_html,
                      reply_to_id=None, attachments_data=None, internet_message_id=None,
                      deal_id=None, alert_type=None, db=None):
    """Send email via Graph API.

    Note: internet_message_id is accepted for backward compatibility but
    Graph API rejects standard In-Reply-To/References headers (must start with x-).
    For proper threading, use helper_graph_reply() with the Graph message ID instead.

    Optional gate params (deal_id, alert_type, db) enable Firestore dedup checks.
    """
    # ── HARD BLOCK: Never send to external recipients ──
    if not _is_internal_recipient(to_email):
        print(f"\U0001f6ab BLOCKED outbound to external: {to_email} | subj={(subject or '')[:60]}")
        try:
            _db = db
            if _db is None:
                # Try to get db from caller context — but don't crash if unavailable
                pass
            if _db is not None:
                _db.collection("security_log").add({
                    "type": "BLOCKED_EXTERNAL_SEND",
                    "function": "helper_graph_send",
                    "to_email": to_email,
                    "subject": (subject or "")[:120],
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                })
        except Exception:
            pass  # logging failure never blocks the guard
        return False

    # ── Email quality gate (fail-open) ──
    try:
        approved, reason = email_quality_gate(
            to_email, subject, body_html,
            deal_id=deal_id, alert_type=alert_type, db=db)
        if not approved:
            print(f"\U0001f4e7 Email BLOCKED by quality gate: {reason} | to={to_email} subj={(subject or '')[:60]}")
            return False
    except Exception:
        pass  # fail open — gate error never blocks sending

    try:
        url = f"https://graph.microsoft.com/v1.0/users/{user_email}/sendMail"
        message = {
            'subject': f"Re: {subject}" if reply_to_id else subject,
            'body': {'contentType': 'HTML', 'content': body_html},
            'toRecipients': [{'emailAddress': {'address': to_email}}]
        }
        if attachments_data:
            message['attachments'] = [
                {
                    '@odata.type': '#microsoft.graph.fileAttachment',
                    'name': a.get('name', 'file'),
                    'contentType': a.get('contentType', 'application/octet-stream'),
                    'contentBytes': a.get('contentBytes')
                } for a in attachments_data if a.get('contentBytes')
            ]
        response = requests.post(url, headers={'Authorization': f'Bearer {access_token}', 'Content-Type': 'application/json'}, json={'message': message, 'saveToSentItems': True})
        return response.status_code == 202
    except Exception as e:
        print(f"Send error: {e}")
        return False


def helper_graph_reply(access_token, user_email, message_id, body_html, to_email=None, cc_emails=None, subject=None, db=None, deal_id=None, alert_type=None):
    """Reply to an existing email via Graph API (threads correctly in Outlook).

    Args:
        subject: Optional override for reply subject line. If provided,
                 replaces the default "Re: ..." subject inherited from the thread.
        db: Optional Firestore client for security_log logging.
        deal_id: Optional deal ID for dedup checks via email_quality_gate.
        alert_type: Optional alert type for dedup checks via email_quality_gate.
    """
    # ── HARD BLOCK: Never reply to external recipients ──
    if not to_email:
        print(f"\U0001f6ab BLOCKED reply: to_email is None (cannot verify recipient)")
        try:
            if db is not None:
                db.collection("security_log").add({
                    "type": "BLOCKED_EXTERNAL_REPLY",
                    "function": "helper_graph_reply",
                    "reason": "to_email_is_none",
                    "subject": (subject or "")[:120],
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                })
        except Exception:
            pass
        return False
    if not _is_internal_recipient(to_email):
        print(f"\U0001f6ab BLOCKED reply to external: {to_email}")
        try:
            if db is not None:
                db.collection("security_log").add({
                    "type": "BLOCKED_EXTERNAL_REPLY",
                    "function": "helper_graph_reply",
                    "to_email": to_email,
                    "subject": (subject or "")[:120],
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                })
        except Exception:
            pass
        return False
    if cc_emails:
        cc_emails = [e for e in cc_emails if _is_internal_recipient(e)]

    # ── Email quality gate (fail-open) ──
    try:
        approved, reason = email_quality_gate(
            to_email, subject or "", body_html,
            deal_id=deal_id, alert_type=alert_type, db=db)
        if not approved:
            print(f"\U0001f4e7 Reply BLOCKED by quality gate: {reason} | to={to_email}")
            return False
    except Exception:
        pass  # fail open — gate error never blocks sending

    try:
        url = f"https://graph.microsoft.com/v1.0/users/{user_email}/messages/{message_id}/reply"
        payload = {
            'message': {
                'body': {'contentType': 'HTML', 'content': body_html}
            }
        }
        if subject:
            payload['message']['subject'] = subject
        if to_email:
            payload['message']['toRecipients'] = [{'emailAddress': {'address': to_email}}]
        if cc_emails:
            payload['message']['ccRecipients'] = [
                {'emailAddress': {'address': e}} for e in cc_emails if e
            ]
        response = requests.post(
            url,
            headers={'Authorization': f'Bearer {access_token}', 'Content-Type': 'application/json'},
            json=payload
        )
        return response.status_code == 202
    except Exception as e:
        print(f"Reply error: {e}")
        return False


# ============================================================
# EMAIL QUALITY GUARDS
# ============================================================

_RE_FWD_PATTERN = re.compile(r'^(?:\s*(?:Re|RE|re|Fwd|FWD|FW|Fw|fw)\s*:\s*)+')


def is_direct_recipient(msg, rcb_email):
    """Check if rcb_email is in the message's toRecipients (not just CC).

    Returns True if rcb@ is a direct TO recipient — replies are allowed.
    Returns False if rcb@ is only CC'd — replies must be suppressed.
    Returns True if toRecipients is missing/empty (fail-open for safety).
    """
    to_recipients = msg.get('toRecipients', [])
    if not to_recipients:
        return True  # fail-open: if Graph API didn't return toRecipients, allow
    return any(
        rcb_email.lower() in r.get('emailAddress', {}).get('address', '').lower()
        for r in to_recipients
    )



def clean_email_subject(subject):
    """Strip accumulated Re:/RE:/Fwd:/FW: prefix chains from a subject line."""
    if not subject:
        return ""
    cleaned = _RE_FWD_PATTERN.sub('', subject).strip()
    return cleaned if cleaned else subject


def validate_email_before_send(subject, body_html):
    """Check email has meaningful subject and body content.

    Returns (is_valid, reason) tuple.
    Rejects: empty subject, garbage subject, empty body, body with only template dashes.
    """
    # Check subject
    if not subject or not subject.strip():
        return False, "empty_subject"
    # Strip Re:/Fwd: directly (clean_email_subject returns original if all-prefix)
    stripped_subj = _RE_FWD_PATTERN.sub('', subject).strip()
    if not stripped_subj or len(stripped_subj) < 3:
        return False, "garbage_subject"
    # Check body
    if not body_html or not body_html.strip():
        return False, "empty_body"
    # Strip HTML tags and entities, check remaining text
    text_only = re.sub(r'<[^>]+>', '', body_html)
    text_only = re.sub(r'&[a-zA-Z]+;|&#\d+;', ' ', text_only)
    text_only = text_only.strip()
    if len(text_only) < 20:
        return False, "body_too_short"
    # Reject templates that are all dashes / em-dashes / whitespace
    meaningful = re.sub(r'[\s\-\u2014\u2013]+', '', text_only)
    if len(meaningful) < 10:
        return False, "body_only_dashes"
    return True, "ok"


# ============================================================
# EMAIL QUALITY GATE (Session 45)
# ============================================================

_GENERIC_SUBJECTS = frozenset({
    "test", "hello", "hi", "שלום", "בדיקה", "no subject",
    "notification", "update", "fyi", "info",
})

_PLACEHOLDER_RE = re.compile(
    r'^[\s\-\u2014\u2013]*$'
    r'|^טרם הגיע$'
    r'|^—$'
    r'|^N/?A$',
    re.IGNORECASE,
)

_TD_RE = re.compile(r'<td[^>]*>(.*?)</td>', re.DOTALL | re.IGNORECASE)
_HTML_TAG_RE = re.compile(r'<[^>]+>')


def _log_email_quality(db, approved, reason, recipient, subject,
                       deal_id, alert_type, content_hash):
    """Log gate decision to Firestore. Fire-and-forget, never raises."""
    if not db:
        return
    try:
        now = datetime.now(timezone.utc)
        now_iso = now.isoformat()
        record = {
            "approved": approved,
            "reason": reason,
            "recipient": recipient or "",
            "subject": (subject or "")[:200],
            "deal_id": deal_id or "",
            "alert_type": alert_type or "",
            "content_hash": content_hash or "",
            "timestamp": now_iso,
        }
        db.collection("email_quality_log").add(record)

        # Update dedup tracking docs on approval only
        if approved:
            if deal_id and alert_type:
                dedup_key = f"dedup_{deal_id}_{alert_type}"
                db.collection("email_quality_log").document(dedup_key).set({
                    "timestamp": now_iso,
                    "recipient": recipient or "",
                    "type": "dedup_tracker",
                })
            if deal_id and recipient:
                subj_key = f"subj_{deal_id}_{hashlib.md5((recipient or '').lower().encode()).hexdigest()[:8]}"
                db.collection("email_quality_log").document(subj_key).set({
                    "subject": subject or "",
                    "timestamp": now_iso,
                    "type": "subject_tracker",
                })
            is_digest = (
                "digest" in (alert_type or "").lower()
                or "\u05d3\u05d5\u05d7 \u05d1\u05d5\u05e7\u05e8" in (subject or "")
                or "digest" in (subject or "").lower()
            )
            if is_digest and recipient:
                digest_key = f"digest_{hashlib.md5((recipient or '').lower().encode()).hexdigest()[:8]}"
                db.collection("email_quality_log").document(digest_key).set({
                    "content_hash": content_hash or "",
                    "timestamp": now_iso,
                    "type": "digest_tracker",
                })
    except Exception as e:
        print(f"\u26a0\ufe0f _log_email_quality error (non-fatal): {e}")


def email_quality_gate(recipient, subject, html_body,
                       deal_id=None, alert_type=None, db=None):
    """Central email quality gate. Every outgoing email passes through here.

    Returns (approved: bool, reason: str).
    Fail-open: any internal error returns (True, "gate_error_failopen").
    """
    try:
        # ── Rule 5: never send to self ──
        if recipient and recipient.lower().strip() == "rcb@rpa-port.co.il":
            _log_email_quality(db, False, "self_send", recipient, subject,
                               deal_id, alert_type, None)
            return False, "self_send"

        # ── Rule 1: body empty or under 200 characters ──
        if not html_body or len(html_body) < 200:
            _log_email_quality(db, False, "body_under_200", recipient, subject,
                               deal_id, alert_type, None)
            return False, "body_under_200"

        # ── Rule 2a: subject empty or generic ──
        if not subject or not subject.strip():
            _log_email_quality(db, False, "empty_subject", recipient, subject,
                               deal_id, alert_type, None)
            return False, "empty_subject"
        stripped_subj = _RE_FWD_PATTERN.sub('', subject).strip()
        if not stripped_subj or len(stripped_subj) < 3:
            _log_email_quality(db, False, "generic_subject", recipient, subject,
                               deal_id, alert_type, None)
            return False, "generic_subject"
        if stripped_subj.lower() in _GENERIC_SUBJECTS:
            _log_email_quality(db, False, "generic_subject", recipient, subject,
                               deal_id, alert_type, None)
            return False, "generic_subject"

        # ── Rule 3: all data cells are placeholder ──
        td_cells = _TD_RE.findall(html_body)
        if len(td_cells) > 3:
            real_data = 0
            for cell_html in td_cells:
                cell_text = _HTML_TAG_RE.sub('', cell_html).strip()
                if cell_text and not _PLACEHOLDER_RE.match(cell_text):
                    real_data += 1
            if real_data == 0:
                _log_email_quality(db, False, "all_placeholder_data", recipient,
                                   subject, deal_id, alert_type, None)
                return False, "all_placeholder_data"

        # ── Rule 7 (Session 47): block empty classification emails ──
        # A classification email (subject contains "סיווג") with zero HS codes
        # means extraction failed — this should not reach the user.
        if "\u05e1\u05d9\u05d5\u05d5\u05d2" in (subject or "") and "RCB-" in (subject or ""):
            _hs_cells = re.findall(r'\d{2}\.\d{2}\.\d{4,6}', html_body or "")
            if not _hs_cells:
                _log_email_quality(db, False, "empty_classification", recipient,
                                   subject, deal_id, alert_type, None)
                return False, "empty_classification"

        # ── Content hash for dedup ──
        content_hash = hashlib.md5(
            html_body.encode('utf-8', errors='replace')
        ).hexdigest()

        # ── Firestore-dependent rules (skip if no db) ──
        if db:
            try:
                # Rule 4: same deal_id + alert_type within 4 hours
                if deal_id and alert_type:
                    dedup_key = f"dedup_{deal_id}_{alert_type}"
                    dedup_doc = db.collection("email_quality_log").document(dedup_key).get()
                    if dedup_doc.exists:
                        last_ts_str = dedup_doc.to_dict().get("timestamp", "")
                        if last_ts_str:
                            last_ts = datetime.fromisoformat(last_ts_str)
                            if not last_ts.tzinfo:
                                last_ts = last_ts.replace(tzinfo=timezone.utc)
                            age = datetime.now(timezone.utc) - last_ts
                            if age.total_seconds() < 4 * 3600:
                                _log_email_quality(db, False, "dedup_4h", recipient,
                                                   subject, deal_id, alert_type,
                                                   content_hash)
                                return False, "dedup_4h"

                # Rule 2b: subject unchanged from last send to same deal+recipient
                if deal_id and recipient:
                    subj_key = f"subj_{deal_id}_{hashlib.md5((recipient or '').lower().encode()).hexdigest()[:8]}"
                    subj_doc = db.collection("email_quality_log").document(subj_key).get()
                    if subj_doc.exists:
                        if subj_doc.to_dict().get("subject") == subject:
                            _log_email_quality(db, False, "unchanged_subject",
                                               recipient, subject, deal_id,
                                               alert_type, content_hash)
                            return False, "unchanged_subject"

                # Rule 6: digest content identical to last digest
                is_digest = (
                    "digest" in (alert_type or "").lower()
                    or "\u05d3\u05d5\u05d7 \u05d1\u05d5\u05e7\u05e8" in (subject or "")
                    or "digest" in (subject or "").lower()
                )
                if is_digest and recipient:
                    digest_key = f"digest_{hashlib.md5((recipient or '').lower().encode()).hexdigest()[:8]}"
                    digest_doc = db.collection("email_quality_log").document(digest_key).get()
                    if digest_doc.exists:
                        if digest_doc.to_dict().get("content_hash") == content_hash:
                            _log_email_quality(db, False, "duplicate_digest",
                                               recipient, subject, deal_id,
                                               alert_type, content_hash)
                            return False, "duplicate_digest"
            except Exception as e:
                # Firestore errors → fail open, skip remaining Firestore rules
                print(f"\u26a0\ufe0f email_quality_gate Firestore error (skipping): {e}")

        # ── Approved ──
        _log_email_quality(db, True, "approved", recipient, subject,
                           deal_id, alert_type, content_hash)
        return True, "approved"

    except Exception as e:
        # Fail open — never block legitimate email
        print(f"\u26a0\ufe0f email_quality_gate error (fail-open): {e}")
        return True, "gate_error_failopen"


# ============================================================
# HEBREW NAMES
# ============================================================

HEBREW_NAMES = {
    'doron': 'דורון', 'amit': 'עמית', 'yossi': 'יוסי', 'moshe': 'משה', 
    'david': 'דוד', 'avi': 'אבי', 'eli': 'אלי', 'dan': 'דן', 
    'oren': 'אורן', 'guy': 'גיא', 'tal': 'טל', 'nir': 'ניר', 
    'roy': 'רועי', 'gal': 'גל', 'yuval': 'יובל', 'itay': 'איתי', 
    'idan': 'עידן', 'eyal': 'אייל', 'ran': 'רן', 'uri': 'אורי', 
    'yair': 'יאיר', 'alon': 'אלון', 'chen': 'חן', 'lior': 'ליאור', 
    'omer': 'עומר', 'noam': 'נועם', 'ron': 'רון', 'yoni': 'יוני',
    'asaf': 'אסף', 'matan': 'מתן', 'nadav': 'נדב', 'raz': 'רז'
}

def to_hebrew_name(name):
    """Convert English name to Hebrew if possible"""
    if not name:
        return "שלום"
    if any('\u0590' <= c <= '\u05FF' for c in name):
        return name
    return HEBREW_NAMES.get(name.lower().strip(), name)

# ============================================================
# EMAIL BUILDERS
# ============================================================

def build_rcb_reply(sender_name, attachments, subject="", is_first_email=False, include_joke=False):
    """Build acknowledgment reply (Email 1) — table-based, Outlook-safe, RTL."""
    name = sender_name.split('<')[0].strip()
    name = to_hebrew_name(name.split()[0]) if name and '@' not in name else "שלום"

    hour = datetime.now().hour
    greeting = "בוקר טוב" if 5 <= hour < 12 else "צהריים טובים" if 12 <= hour < 17 else "ערב טוב" if 17 <= hour < 21 else "לילה טוב"

    # ── Attachments list (table rows, not <ul>) ──
    att_rows = ""
    if attachments:
        for a in attachments:
            ext = a.get('type', '')
            icon = "&#128209;" if ext == '.pdf' else "&#128202;" if ext in ['.xlsx', '.xls'] else "&#128444;" if ext in ['.jpg', '.png'] else "&#128196;"
            fname = a.get('filename', '\u05e7\u05d5\u05d1\u05e5')
            att_rows += (
                f'<tr><td style="padding:4px 10px;font-family:Arial,sans-serif;font-size:13px;'
                f'color:#333333;border-bottom:1px solid #f0f0f0;">{icon} {fname}</td></tr>'
            )

    att_section = ""
    if att_rows:
        att_section = (
            '<tr><td style="padding:10px 30px 0 30px;">'
            '<table width="100%" cellpadding="0" cellspacing="0" dir="rtl" '
            'style="background:#f8f9fa;border:1px solid #dee2e6;border-radius:6px;">'
            '<tr><td style="padding:10px;font-family:Arial,sans-serif;font-size:13px;'
            f'font-weight:bold;color:#1e3a5f;border-bottom:1px solid #dee2e6;">'
            f'&#128269; \u05d6\u05d9\u05d4\u05d9\u05ea\u05d9 {len(attachments)} \u05e7\u05d1\u05e6\u05d9\u05dd:</td></tr>'
            f'{att_rows}'
            '</table></td></tr>'
        )

    subject_row = ""
    if subject:
        subject_row = (
            '<tr><td style="padding:10px 30px 0 30px;font-family:Arial,sans-serif;'
            f'font-size:13px;color:#333333;">&#128203; <strong>\u05e0\u05d5\u05e9\u05d0:</strong> {subject}</td></tr>'
        )

    html = (
        '<!DOCTYPE html>\n'
        '<html dir="rtl" lang="he">\n'
        '<head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1.0"></head>\n'
        '<body dir="rtl" style="margin:0;padding:0;background:#f5f5f5;font-family:Arial,Helvetica,sans-serif;">\n'
        '<table width="100%" cellpadding="0" cellspacing="0" style="background:#f5f5f5;padding:20px 0;">\n'
        '<tr><td align="center">\n'
        '<table width="640" cellpadding="0" cellspacing="0" style="max-width:640px;width:100%;'
        'background:#ffffff;border-radius:8px;overflow:hidden;box-shadow:0 2px 8px rgba(0,0,0,0.1);">\n'
        # ── Header bar ──
        '<tr><td style="background:#1e3a5f;padding:20px 30px;">'
        '<table width="100%" cellpadding="0" cellspacing="0"><tr>'
        '<td valign="middle">'
        '<img src="https://rpa-port.com/wp-content/uploads/2016/09/logo.png" alt="RPA-PORT" '
        'width="48" height="48" style="display:inline-block;vertical-align:middle;border:0;">'
        '<span style="display:inline-block;vertical-align:middle;padding-left:12px;">'
        '<span style="color:#ffffff;font-size:18px;font-weight:bold;">R.P.A. PORT LTD</span><br>'
        '<span style="color:#aed6f1;font-size:13px;">RCB &mdash; AI Customs Broker</span>'
        '</span></td>'
        '</tr></table></td></tr>\n'
        # ── Greeting ──
        '<tr><td style="padding:25px 30px 10px 30px;font-family:Arial,sans-serif;font-size:15px;'
        f'color:#333333;line-height:1.6;"><strong>{greeting} {name},</strong></td></tr>\n'
        # ── Body text ──
        '<tr><td style="padding:0 30px 10px 30px;font-family:Arial,sans-serif;font-size:14px;'
        'color:#333333;line-height:1.6;">'
        '\u05ea\u05d5\u05d3\u05d4 \u05e2\u05dc \u05e4\u05e0\u05d9\u05d9\u05ea\u05da! '
        '\u05e7\u05d9\u05d1\u05dc\u05ea\u05d9 \u05d0\u05ea \u05d4\u05de\u05e1\u05de\u05db\u05d9\u05dd '
        '\u05e9\u05dc\u05da \u05d5\u05d0\u05e0\u05d9 \u05de\u05ea\u05d7\u05d9\u05dc\u05d4 '
        '\u05dc\u05e2\u05d1\u05d5\u05e8 \u05e2\u05dc\u05d9\u05d4\u05dd.</td></tr>\n'
        f'{subject_row}'
        f'{att_section}'
        # ── Status note ──
        '<tr><td style="padding:15px 30px;font-family:Arial,sans-serif;font-size:14px;'
        'color:#333333;line-height:1.6;">'
        '&#9203; \u05d0\u05e2\u05d1\u05d5\u05e8 \u05e2\u05dc \u05d4\u05de\u05e1\u05de\u05db\u05d9\u05dd '
        '\u05d5\u05d0\u05e9\u05dc\u05d7 \u05dc\u05da \u05d3\u05d5"\u05d7 \u05e1\u05d9\u05d5\u05d5\u05d2 '
        '\u05de\u05db\u05e1 \u05de\u05e4\u05d5\u05e8\u05d8 \u05ea\u05d5\u05da \u05de\u05e1\u05e4\u05e8 '
        '\u05d3\u05e7\u05d5\u05ea.</td></tr>\n'
        # ── Divider ──
        '<tr><td style="padding:0 30px;">'
        '<table width="100%" cellpadding="0" cellspacing="0">'
        '<tr><td style="border-top:1px solid #dee2e6;font-size:1px;height:1px;">&nbsp;</td></tr>'
        '</table></td></tr>\n'
        # ── Signature ──
        '<tr><td style="padding:15px 30px;">'
        '<table dir="rtl" cellpadding="0" cellspacing="0"><tr>'
        '<td style="padding-left:15px;" valign="top">'
        '<img src="https://rpa-port.com/wp-content/uploads/2016/09/logo.png" '
        'width="48" height="48" style="border:0;"></td>'
        '<td style="border-right:3px solid #1e3a5f;padding-right:15px;" valign="top">'
        '<strong style="color:#1e3a5f;font-family:Arial,sans-serif;font-size:13px;">'
        '&#129302; RCB - AI Customs Broker</strong><br>'
        '<strong style="font-family:Arial,sans-serif;font-size:12px;">R.P.A. PORT LTD</strong><br>'
        '<span style="font-family:Arial,sans-serif;font-size:11px;color:#666666;">'
        '&#128231; rcb@rpa-port.co.il</span>'
        '</td></tr></table></td></tr>\n'
        # ── Close ──
        '</table>\n</td></tr></table>\n'
        '</body></html>'
    )

    return html

def get_anthropic_key(get_secret_func):
    """Get Anthropic API key"""
    try:
        return get_secret_func('ANTHROPIC_API_KEY')
    except Exception as e:
        print(f"Anthropic key error: {e}")
        return None
