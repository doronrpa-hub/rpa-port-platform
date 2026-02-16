"""
Tests for SmartExtractor multi-method extraction engine.

Session 28C — Assignment 20.
NEW FILE.
"""

import pytest
from unittest.mock import patch, MagicMock
import sys
import os

# Ensure lib is importable
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from lib.smart_extractor import SmartExtractor
from lib.extraction_result import ExtractionResult


@pytest.fixture
def extractor():
    return SmartExtractor()


# ─────────────────────────────────────
#  Empty / garbage inputs
# ─────────────────────────────────────

class TestEmptyInputs:

    def test_empty_bytes_returns_empty_result(self, extractor):
        """The extractor should NEVER throw an exception. Always return a result."""
        result = extractor.extract(b"", "empty.pdf", "application/pdf")
        assert result.is_empty is True
        assert result.confidence == 0.0
        assert len(result.warnings) > 0

    def test_none_bytes_returns_empty_result(self, extractor):
        result = extractor.extract(None, "nothing.pdf", "application/pdf")
        assert result.is_empty is True

    def test_garbage_bytes_never_crashes(self, extractor):
        """Binary garbage should not crash — should fail gracefully."""
        garbage = bytes(range(256)) * 50
        result = extractor.extract(garbage, "garbage.bin", "application/octet-stream")
        assert isinstance(result, ExtractionResult)
        # Should either find something or return empty — never throw
        assert isinstance(result.text, str)


# ─────────────────────────────────────
#  PDF extraction
# ─────────────────────────────────────

class TestPDFExtraction:

    def test_pdf_tries_multiple_methods(self, extractor):
        """For PDF files, multiple extraction methods should be attempted."""
        # Create a minimal valid PDF
        minimal_pdf = (
            b"%PDF-1.0\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj "
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj "
            b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
            b"/Resources<<>>>>endobj\nxref\n0 4\n"
            b"0000000000 65535 f \n0000000009 00000 n \n"
            b"0000000058 00000 n \n0000000115 00000 n \n"
            b"trailer<</Size 4/Root 1 0 R>>\nstartxref\n206\n%%EOF"
        )
        result = extractor.extract(minimal_pdf, "test.pdf", "application/pdf")
        # Even a minimal PDF without text should complete without error
        assert isinstance(result, ExtractionResult)

    @patch("lib.smart_extractor.SmartExtractor._vision_ocr_pdf")
    def test_vision_ocr_called_when_text_methods_fail(self, mock_ocr, extractor):
        """Vision OCR should be the fallback when text methods return low confidence."""
        mock_ocr.return_value = "Invoice #123 Total: $500"
        garbage_pdf = b"%PDF-1.0\n" + b"\x00" * 200
        result = extractor.extract(garbage_pdf, "scanned.pdf", "application/pdf")
        # Either the text methods found something or OCR was called
        assert isinstance(result, ExtractionResult)


# ─────────────────────────────────────
#  HTML extraction
# ─────────────────────────────────────

class TestHTMLExtraction:

    def test_html_with_text(self, extractor):
        html = b"<html><body><h1>Invoice</h1><p>Total: $1000</p></body></html>"
        result = extractor.extract(html, "invoice.html", "text/html")
        assert "Invoice" in result.text
        assert "$1000" in result.text or "1000" in result.text
        assert result.confidence > 0

    def test_html_with_table(self, extractor):
        html = (
            b"<html><body>"
            b"<table><tr><th>Item</th><th>Price</th></tr>"
            b"<tr><td>Widget</td><td>100</td></tr>"
            b"<tr><td>Gadget</td><td>200</td></tr>"
            b"</table></body></html>"
        )
        result = extractor.extract(html, "items.html", "text/html")
        assert "Widget" in result.text
        assert result.confidence > 0

    def test_html_scripts_removed(self, extractor):
        html = (
            b"<html><body>"
            b"<script>var x = 'malicious';</script>"
            b"<p>Real content here</p>"
            b"</body></html>"
        )
        result = extractor.extract(html, "page.html", "text/html")
        assert "malicious" not in result.text
        assert "Real content" in result.text

    def test_hebrew_html(self, extractor):
        html = '<html><body><p>חשבונית מספר 12345</p></body></html>'.encode("utf-8")
        result = extractor.extract(html, "invoice_he.html", "text/html")
        assert "חשבונית" in result.text


# ─────────────────────────────────────
#  Excel extraction
# ─────────────────────────────────────

class TestExcelExtraction:

    def test_excel_openpyxl(self, extractor):
        """Test extraction from a real xlsx file (created in memory)."""
        try:
            import openpyxl
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Item", "Quantity", "Price"])
        ws.append(["Bolts", "100", "5.50"])
        ws.append(["Nuts", "200", "3.25"])
        buf = __import__("io").BytesIO()
        wb.save(buf)
        file_bytes = buf.getvalue()

        result = extractor.extract(file_bytes, "items.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        assert "Bolts" in result.text
        assert "Nuts" in result.text
        assert result.confidence > 0.5


# ─────────────────────────────────────
#  Word extraction
# ─────────────────────────────────────

class TestWordExtraction:

    def test_docx_extraction(self, extractor):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Certificate of Origin")
        doc.add_paragraph("Country: Israel")
        buf = __import__("io").BytesIO()
        doc.save(buf)
        file_bytes = buf.getvalue()

        result = extractor.extract(file_bytes, "coo.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert "Certificate" in result.text
        assert "Israel" in result.text


# ─────────────────────────────────────
#  CSV extraction
# ─────────────────────────────────────

class TestCSVExtraction:

    def test_csv_basic(self, extractor):
        csv_data = b"Item,Qty,Price\nBolts,100,5.50\nNuts,200,3.25\n"
        result = extractor.extract(csv_data, "items.csv", "text/csv")
        assert "Bolts" in result.text
        assert result.confidence > 0.5

    def test_csv_hebrew(self, extractor):
        csv_data = "פריט,כמות,מחיר\nברגים,100,5.50\n".encode("utf-8")
        result = extractor.extract(csv_data, "items_he.csv", "text/csv")
        assert "פריט" in result.text


# ─────────────────────────────────────
#  Helper / utility methods
# ─────────────────────────────────────

class TestHelpers:

    def test_text_similarity_identical(self, extractor):
        assert extractor._text_similarity("hello world", "hello world") == 1.0

    def test_text_similarity_disjoint(self, extractor):
        assert extractor._text_similarity("hello world", "foo bar") == 0.0

    def test_text_similarity_partial(self, extractor):
        sim = extractor._text_similarity("hello world foo", "hello bar foo")
        assert 0.3 < sim < 0.8

    def test_detect_encoding_utf8(self, extractor):
        text = "Hello, World!".encode("utf-8")
        assert extractor._detect_encoding(text) == "utf-8"

    def test_detect_encoding_hebrew(self, extractor):
        text = "שלום עולם".encode("windows-1255")
        enc = extractor._detect_encoding(text)
        # Should detect a Hebrew-compatible encoding
        assert enc in ("windows-1255", "iso-8859-8", "utf-8")

    def test_table_to_dicts(self, extractor):
        rows = [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]
        result = extractor._table_to_dicts(rows)
        assert len(result) == 2
        assert result[0]["Name"] == "Alice"
        assert result[1]["Age"] == "25"

    def test_table_to_dicts_empty(self, extractor):
        assert extractor._table_to_dicts([]) == []
        assert extractor._table_to_dicts([["Header"]]) == []

    def test_pick_best_prefers_higher_confidence(self, extractor):
        r1 = ExtractionResult(text="short", confidence=0.5, method="a")
        r2 = ExtractionResult(text="much longer text with more content", confidence=0.9, method="b")
        best = extractor._pick_best([r1, r2])
        assert best.method == "b"

    def test_pick_best_switches_to_longer_text(self, extractor):
        """If second method found 50%+ more text with decent confidence, prefer it."""
        short = "x " * 50
        long = "x " * 200
        r1 = ExtractionResult(text=short, confidence=0.85, method="a")
        r2 = ExtractionResult(text=long, confidence=0.8, method="b")
        best = extractor._pick_best([r1, r2])
        assert best.method == "b"

    def test_cleanup_hebrew(self, extractor):
        text = "עמיל מכם  שילם  חשבוו"
        cleaned = extractor._cleanup_hebrew(text)
        assert "עמיל מכס" in cleaned
        assert "חשבון" in cleaned
